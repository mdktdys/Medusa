import datetime
from typing import List

from docx import Document
from docx.table import Table
import numpy as np

from src.parser.schemas.paras import Paras
from src.parser.models.data_model import Data
from src.parser.models.course_model import Course
from src.parser.models.group_model import Group
from src.parser.models.teacher_model import Teacher
from src.parser.models.cabinet_model import Cabinet
from src.parser.parsers import init_date_model
from src.parser.supabase import SupaBaseWorker
from src.parser.shared import (
    get_align_course_by_group,
    get_teacher_from_string,
    get_cabinet_from_string,
    get_group_from_string,
)

def schedule_parser_v4(data_model: Data, monday_date: datetime.date):
    def _extract_all_tables_to_rows(tables: list[Table]) -> list[list[str]]:
        rows = []
        for table in tables:
            for row in table.rows:
                data = []
                try:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell.tables:
                            nested_table_rows = _extract_all_tables_to_rows(cell.tables)
                            data.extend(nested_table_rows)
                        else:
                            data.append(cell_text)
                    rows.append(data)
                except Exception as e:
                    print(e)
        return rows

    docx: Document = Document('../Общее-расписание-с-01.04 (1).docx')
    tables = docx.tables
    paras: List[Paras] = []
    errors: List[str] = []

    for tableIndex in range(len(tables)):
        rows = _extract_all_tables_to_rows([tables[tableIndex]])
        rows.pop(1)

        # for row in rows:
        #     row.pop(0)
        #     row.pop(0)

        groups_count = len(rows[0]) / 3


        for group_index in range(int(groups_count)):
            group_name = rows[0][0 + (group_index * 3)]
            print(group_name)

            for day_index in range(6):
                for para_index in range(7):
                    course_name = rows[1 + (day_index * 7) + para_index][0 + (group_index * 3)]
                    teacher_name = rows[1 + (day_index * 7) + para_index][1 + (group_index * 3)]
                    cabinet_name = rows[1 + (day_index * 7) + para_index][2 + (group_index * 3)]

                    if course_name is np.nan or course_name == '':
                        continue

                    # print(group_name, course_name, teacher_name, cabinet_name, para_index + 1, day_index)

                    group: Group = get_group_from_string(group_name, data_model.GROUPS)
                    course: Course | None = get_align_course_by_group(group, course_name, data_model)

                    if course is None:
                        errors.append(f"Course not found {course_name} in {group_name}({group.id})")
                        continue

                    teacher: Teacher | None = get_teacher_from_string(teacher_name, data_model.TEACHERS)

                    if teacher is None:
                        errors.append(f"Teacher not found {teacher_name} in {group_name}")
                        continue

                    cabinet: Cabinet = get_cabinet_from_string(cabinet_name, data_model.CABINETS)

                    if cabinet is None:
                        errors.append(f"Cabinet not found {cabinet_name} in {group_name}")
                        continue

                    date: datetime.date = monday_date + datetime.timedelta(days=day_index)

                    paras.append(Paras(group=group.id, number=para_index + 1, course=course.id, teacher=teacher.id, cabinet=cabinet.id, date=date.strftime("%Y-%m-%d")))

    if len(errors) > 0:
        raise Exception(errors)

    return paras


if __name__ == '__main__':
    supabase_client = SupaBaseWorker()
    data_model = init_date_model(supabase_client)
    rows = schedule_parser_v4(data_model=data_model, monday_date=datetime.date(2025,3,31))
    supabase_paras = []

    for para in rows:
        print(para)
        supabase_paras.append(
            {
                "group": para.group,
                "number": para.number,
                "course": para.course,
                "teacher": para.teacher,
                "cabinet": para.cabinet,
                "date": para.date,
            }
        )

    res = supabase_client.client.table("Paras").insert(supabase_paras).execute()
    print(rows)