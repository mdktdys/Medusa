"""
Модуль содержит в себе парсер замен и вспомогательные методы для него
"""

import asyncio
import base64
import hashlib
import os
import re
from io import BytesIO
from typing import List, Tuple
import fitz
import requests
from docx import Document
from src.utils.ai_requests import send_ai_request, extract_groups_promt, teacher_cabinets_switches_promt
from docx.table import Table
from datetime import date
import aspose.words as aw
from src.parser.models.cabinet_model import Cabinet
from src.parser.models.course_model import Course
from src.parser.models.data_model import Data
from src.parser.models.group_model import Group
from src.parser.models.teacher_model import Teacher
from src.parser.schemas.parse_zamena_schemas import (
    ZamenaParseFailedNotFoundItems,
    ZamenaParseResult,
    ZamenaParseResultJson,
    ZamenaParseSucess,
)
from src.parser.shared import (
    get_align_course_by_group,
    get_teacher_from_string,
    clean_dirty_string,
    get_empty_course,
    is_empty_course,
    get_cabinet_from_string,
    get_group_from_string,
)
from src.parser.supabase import SupaBaseWorker

not_found_items: List[str] = []
parse_result = None


def parse_zamena_v2(stream, data_model, link, date: date, supabase_client) -> ZamenaParseResult:
    all_rows, header_paragraphs = _get_all_tables(stream)
    header: str = ' '.join(head.text for head in header_paragraphs if 'Исп.' not in head.text)

    practice_groups: List[Group] = _extract_practice_groups(header, data_model)
    practice_groups_: list[int] = [i.id for i in practice_groups]
    
    teacher_cabinets_switches: list[Tuple[int, int]] = [(pair[0].id, pair[1].id) for pair in extract_teacher_cabinets_switchers(header, data_model)]

    work_rows = _prepare_work_rows(all_rows)
    work_rows = _filter_and_clean_rows(work_rows)

    work_rows, full_zamena_groups, liquidation = handle_special_cases(work_rows, data_model)
    update_empty_group_column(work_rows)
    work_rows = process_multiple_entries(work_rows)

    map_entities_to_ids(work_rows, data_model)

    if len(not_found_items) > 0:
        return ZamenaParseFailedNotFoundItems(
            error="Not found items",
            items=list(set(not_found_items)),
            result="error",
            trace=f"{link}",
        )

    zamenas = [{"group": i[0], "number": int(i[1]), "course": i[4], "teacher": i[5], "cabinet": i[6]} for i in work_rows]
    full_zamenas_groups: list[int] = [get_group_by_id(target_name=i,data_model=data_model,groups=data_model.GROUPS,supabase_client=supabase_client).id for i in full_zamena_groups]
    hash = get_remote_file_hash(link)
    
    return ZamenaParseResultJson(
        result = 'ok',
        zamenas = zamenas,
        teacher_cabinet_switches = teacher_cabinets_switches,
        liquidation_groups = liquidation,
        full_zamena_groups = full_zamenas_groups,
        practice_groups = practice_groups_,
        file_hash = hash,
        date = date,
    )


def parseZamenas(
        stream: BytesIO,
        date_: date,
        data_model: Data,
        link: str,
        supabase_client: SupaBaseWorker,
        force: bool,
) -> ZamenaParseResult:
    all_rows, header_paragraphs = _get_all_tables(stream)
    header: str = ' '.join(head.text for head in header_paragraphs)

    practice_groups = _extract_practice_groups(header, data_model)
    work_rows = _prepare_work_rows(all_rows)
    work_rows = _filter_and_clean_rows(work_rows)

    work_rows, full_zamena_groups, liquidation = handle_special_cases(work_rows, data_model)
    update_empty_group_column(work_rows)
    work_rows = process_multiple_entries(work_rows)

    map_entities_to_ids(work_rows, data_model)

    if len(not_found_items) > 0:
        return ZamenaParseFailedNotFoundItems(
            error="Not found items",
            items=list(set(not_found_items)),
            result="error",
            trace=f"{link}",
        )

    affected_teachers = get_affected_teachers(work_rows)
    affected_groups = get_affected_groups(work_rows)

    prepare_and_send_supabase_entries(
        work_rows,
        practice_groups,
        liquidation,
        full_zamena_groups,
        date_,
        link,
        data_model,
        supabase_client,
    )

    return ZamenaParseSucess(
        affected_teachers= affected_teachers,
        affected_groups= affected_groups
    )
    
    
def extract_teacher_cabinets_switchers(text: str, data_model: Data) -> List[Tuple[Teacher, Cabinet]]:
    groups_text: str = send_ai_request(request = f'{teacher_cabinets_switches_promt}\n{text}')
    pairs: list[str] = groups_text.split('#')
    extracted: List[Tuple[Teacher, Cabinet]] = []
    
    for pair in pairs:
        separated: list[str] = pair.split('|')
        
        if len(separated) != 2:
            continue
        
        teacher: Teacher | None = get_teacher_from_string(separated[0], teachers = data_model.TEACHERS)
        cabinet: Cabinet | None = get_cabinet_from_string(separated[1], cabinets = data_model.CABINETS)
        
        if teacher is None:
            raise Exception(f'Не найден преподаватель {separated[0]}')
        
        if cabinet is None:
            raise Exception(f'Не найден кабинет {separated[1]}')
        
        extracted.append((teacher, cabinet))
    
    return extracted


def _extract_practice_groups(text: str, data_model: Data):
    groups_text = send_ai_request(request = f'{extract_groups_promt}\n{text}')
    practice_groups: list[Group] = SupaBaseWorker.get_groups_from_string(groups_text, data_model=data_model)
    return practice_groups


def _prepare_work_rows(all_rows):
    """Prepare the rows for processing, flattening nested rows."""
    workRows = []
    for i in all_rows:
        workRows.extend(i if not _is_nested(i) else [i])
    return workRows


def _filter_and_clean_rows(workRows: list[str]):
    """Filter rows based on specific conditions and clean up data."""
    workRows = [i for i in workRows if len(i) == 7]
    for i in workRows:
        if i[2] == "" and i[3] != "":
            i[3] = ""
        if i[1] == "" and i[2] == "" and i[3] == "" and i[5] == "":
            i[4] = ""
    workRows = clearNonDataRows(workRows)
    workRows = clear_empty_sublists(workRows)
    workRows = remove_headers(workRows)
    workRows = removeDemoExam(workRows)
    return workRows


def handle_special_cases(
        workRows: list[str], data_model: Data
) -> Tuple[list, list, list[int]]:
    """Handle specific cases such as liquidation and removing duplicate data."""
    iteration = 0
    liquidation = list()
    fullzamenagroups = list()

    # set group name in zamena group rows
    for i in workRows:
        iteration += 1
        if i[0] == "":
            i[0] = workRows[iteration - 2][0]

    # for i in workRows:
    #     iteration += 1
    #     sample = i[0].strip().lower()
    #     if sample != "" and len(set(i)) == 1:

    # Find full zamena groups
    iteration = 0
    for i in workRows[:]:
        sample = i[0].strip().lower()
        if i[0] != "" and len(set(i)) == 1:
            if "Ликвидация задолженностей".lower().strip() == sample:
                for gr in data_model.GROUPS:
                    if gr.name.lower().strip() in workRows[iteration - 1][0]:
                        liquidation.append(gr.id)
                print(f"REMOVED {i}")
                workRows.remove(i)
                continue
            if "ликвидация" not in i[0].strip().lower():
                fullzamenagroups.append(i[0].strip().replace(" ", "").replace(".", ""))
                workRows.remove(i)
            else:
                try:
                    for gr in data_model.GROUPS:
                        if gr.name.lower().strip() in sample:
                            liquidation.append(gr.id)
                            print(f"Ликвидация {gr.name}")
                except Exception as err:
                    print(err)
                    continue
                workRows.remove(i)

    # for i in workRows[:]:
    #     if i[0] == i[1] and i[2] == i[3]:
    #         if "ликвидация" in i[0].strip().lower():
    #             try:
    #                 sample = i[0].strip().lower()
    #                 for gr in data_model.GROUPS:
    #                     if (gr.name.lower().strip() in sample):
    #                         liquidation.append(gr.id)
    #                         print(f"Ликвидация {gr.name}")
    #             except Exception as err:
    #                 print(err)
    #                 continue
    #             workRows.remove(i)
    #         workRows.remove(i)

    print(20 * "*")
    for i in workRows:
        print(i)

    print(20 * "*")

    # for i in workRows[:]:
    #     i.pop(2)
    #     i.pop(2)
    #     if all(cell == '' for cell in i[1:]):
    #         workRows.remove(i)
    #     elif len(set(i)) == 1:
    #         fullzamenagroups.append(i[0].strip().replace(' ', ""))
    #         workRows.remove(i)

    return workRows, fullzamenagroups, liquidation


def update_empty_group_column(workRows: list[str]):
    """
    Метод удаляет лишние тире, пробелы, запятые, точки и приводит к lowercase
    """
    for i in workRows:
        i[0] = i[0].replace(" ", "").replace(",", "").replace(".", "").lower()
        i[0] = re.sub(r"-{2,}", "-", i[0])


def process_multiple_entries(workRows: list[str]):
    """
    Выполняет обработку над нескольками строк путём удаления лишних запятых, точек и т.п.
    """
    editet = []
    for i in workRows:
        try:
            text = i[1]
            # if text == "12":
            #     text = "1,2"
            text = text.replace(".", ",")
            if text[-1] == ",":
                text = text[:-1]
            if text[0] == ",":
                text = text[1:]
            paras = text.split(",")
            if len(paras) >= 2:
                for para in paras:
                    new = i.copy()
                    new[1] = para
                    editet.append(new)
            else:
                editet.append(i)
        except Exception as error:
            print(error)
            print(i)
            continue
    return editet


def get_affected_groups(work_rows: list) -> List[int]:
    return list(set([i[0] for i in work_rows]))


def get_affected_teachers(work_rows: list) -> List[int]:
    return list(set([i[5] for i in work_rows]))


def map_entities_to_ids(workRows: list, data_model: Data):
    """
    Применяет функцию к Ids и записывает их в строки workRows
    """
    for row in workRows:
        group = get_group_from_string(groups=data_model.GROUPS, string=row[0])
        if clean_dirty_string(row[0]) == clean_dirty_string("423Э-2"):
            group = [group for group in data_model.GROUPS if group.name == "23Э-2"][0]
        if group:
            row[0] = group.id
        else:
            print(f"Not found group in {row[0]}")
            not_found_items.append(f"Not found group in {row[0]}")
        # course = get_course_by_id(
        #     data_model.COURSES, row[2], data_model, supabase_client, args=[group.name]
        # )
        empty_course = get_empty_course(data_model=data_model)

        if is_empty_course(string=row[2], empty_course=empty_course):
            row[2] = empty_course.id
        else:
            course = get_align_course_by_group(
                course_name=row[2], data_model=data_model, group=group
            )
            if course:
                row[2] = course.id
            else:
                not_found_items.append(
                    f"Not found course in {row[2]} group {group.name}"
                )

        teacher = get_teacher_from_string(teachers=data_model.TEACHERS, string=row[3])
        if teacher:
            row[3] = teacher.id
        else:
            not_found_items.append(f"Not found teacher in {row[3]}")
        # teacher = get_teacher_by_id(
        #     data_model.TEACHERS, row[3], data_model, supabase_client
        # )
        # if teacher:
        #     row[3] = teacher.id

        if is_empty_course(string=row[4], empty_course=empty_course):
            row[4] = empty_course.id
        else:
            course = get_align_course_by_group(
                course_name=row[4], data_model=data_model, group=group
            )
            if course:
                row[4] = course.id
            else:
                not_found_items.append(
                    f"Not found course in {row[4]} group {group.name}"
                )

        teacher = get_teacher_from_string(teachers=data_model.TEACHERS, string=row[5])
        if teacher:
            row[5] = teacher.id
        else:
            not_found_items.append(f"Not found teacher in {row[5]}")

        cabinet = get_cabinet_from_string(string=row[6], cabinets=data_model.CABINETS)
        if cabinet:
            row[6] = cabinet.id
        else:
            not_found_items.append(f"Not found cabinet in {row[6]}")


def get_bytes_hash(bytes_data: bytes):
    hasher = hashlib.new("sha256")
    hasher.update(bytes_data)
    return hasher.hexdigest()


def get_remote_file_hash(url):
    response = requests.get(url, stream=True)
    if response.status_code == 200:
        hasher = hashlib.new("sha256")
        for chunk in response.iter_content(chunk_size=1024):
            if chunk:
                hasher.update(chunk)
        return hasher.hexdigest()
    else:
        return None


def get_file_extension(url):
    parts = url.split("/")
    file_name = parts[-1]
    file_parts = file_name.split(".")
    if len(file_parts) > 1:
        return file_parts[-1]
    else:
        return ""


def cleanup_temp_files(file_paths):
    for file_path in file_paths:
        os.remove(file_path)


async def save_pixmap(pixmap, screenshot_path):
    loop = asyncio.get_event_loop()
    await loop.run_in_executor(None, pixmap.save, screenshot_path, "png")


def create_word_screenshots_bytes(word_path) -> List[str]:
    doc = aw.Document(word_path)
    doc.save(f"{word_path}.pdf")
    return create_pdf_screenshots_bytes(f"{word_path}.pdf")


def create_pdf_screenshots_bytes(pdf_path) -> List[str]:
    screenshots_bytes = []
    pdf_document: fitz.Document = fitz.open(f"{pdf_path}.pdf", filetype="pdf")
    for i in range(pdf_document.page_count):
        page: fitz.Page = pdf_document.load_page(i)
        zoom_x = 4  # horizontal zoom
        zoom_y = 4  # vertical zoom
        mat = fitz.Matrix(zoom_x, zoom_y)
        pix: fitz.Pixmap = page.get_pixmap(matrix=mat)
        screenshots_bytes.append(
            base64.b64encode(
                pix.pil_tobytes(format="WEBP", optimize=True, dpi=(400, 400))
            ).decode("utf-8")
        )
    return screenshots_bytes


async def create_pdf_screenshots(pdf_path):
    screenshot_paths = []
    pdf_document: fitz.Document = fitz.open(f"{pdf_path}.pdf")
    for i in range(pdf_document.page_count):
        page: fitz.Page = pdf_document.load_page(i)
        zoom_x = 2  # horizontal zoom
        zoom_y = 2  # vertical zoom
        mat = fitz.Matrix(zoom_x, zoom_y)
        pix: fitz.Pixmap = page.get_pixmap(matrix=mat)
        screenshot_path = f"{pdf_path}_page_{i + 1}.png"
        await save_pixmap(pix, screenshot_path)
        screenshot_paths.append(screenshot_path)
    return screenshot_paths


def download_file(link: str, filename: str):
    response = requests.get(link)
    if response.status_code == 200:
        with open(filename, "wb") as file:
            file.write(response.content)
        print(f"File '{filename}' has been downloaded successfully.")
        return True
    else:
        print("Failed to download the file.")
        return False


def prepare_and_send_supabase_entries(
        workRows,
        practice_groups: list,
        liquidation: list[int],
        fullzamenagroups: list,
        date_: date,
        link,
        data_model: Data,
        supabase_client: SupaBaseWorker,
):
    """
    Подготовить данные и отправить в БД
    """
    practice_supabase = [{"group": i.id, "date": str(date_)} for i in practice_groups]
    zamenas_supabase = [
        {
            "group": i[0],
            "number": int(i[1]),
            "course": i[4],
            "teacher": i[5],
            "cabinet": i[6],
            "date": str(date_),
        }
        for i in workRows
    ]
    full_zamenas_groups = [
        {
            "group": get_group_by_id(
                target_name=i,
                data_model=data_model,
                groups=data_model.GROUPS,
                supabase_client=supabase_client,
            ).id,
            "date": str(date_),
        }
        for i in fullzamenagroups
    ]
    liquidations = [{"group": i, "date": str(date_)} for i in liquidation]

    supabase_client.addZamenas(zamenas=zamenas_supabase)
    if full_zamenas_groups:
        supabase_client.addFullZamenaGroups(groups=full_zamenas_groups)
    if practice_supabase:
        supabase_client.add_practices(practices=practice_supabase)
    if liquidations:
        supabase_client.addLiquidations(liquidations=liquidations)

    hash = get_remote_file_hash(link)
    supabase_client.addNewZamenaFileLink(link, date=date_, hash=hash)


def _get_all_tables(stream: BytesIO):
    docx = Document(stream)
    tables = docx.tables
    all_rows = _extract_all_tables_to_rows(tables)
    return all_rows, docx.paragraphs


def _extract_all_tables_to_rows(tables: list[Table]) -> list[list[str]]:
    rows = []
    for table in tables:
        for row in table.rows:
            data = []
            try:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell.tables:
                        nested_table_rows = _extract_all_tables_to_rows(cell.tables)
                        data.extend(nested_table_rows)
                    else:
                        data.append(cell_text)
                rows.append(data)
            except Exception as e:
                print(e)
    return rows


def _is_nested(row: list[str]) -> bool:
    return isinstance(row, list) and all(isinstance(item, str) for item in row)


def clearNonDataRows(rows: list[list[str]]) -> list[list[str]]:
    filtered_rows = [row for row in rows if len(row) >= 7]
    return filtered_rows


def clear_empty_sublists(nested_list: list[list[str]]) -> list[list[str]]:
    return [sublist for sublist in nested_list if any(item != "" for item in sublist)]


def remove_headers(rows):
    cleared = []
    for i in rows:
        if i[0] != "Группа":
            cleared.append(i)
    return cleared


def removeDemoExam(rows):
    cleared = []
    for i in rows:
        if not i[0].__contains__("экзамен"):
            cleared.append(i)
    return cleared


def find_entity_by_name(entities, target_name, name_key="name", normalize_func=None):
    target_name_normalized = str(target_name).lower()
    if normalize_func:
        target_name_normalized = normalize_func(target_name_normalized)

    for entity in entities:
        entity_name_normalized = getattr(entity, name_key).lower()
        if entity_name_normalized == target_name_normalized:
            return entity
    return None


def add_and_get_entity(entity_type, add_func, entities, target_name, data_model, args):
    entity = find_entity_by_name(entities, target_name)
    if entity:
        return entity
    try:
        not_found_items.append(f"not found {target_name} args: {args}")
        raise Exception(f"not found {target_name}")
        # add_func(target_name, data_model=data_model)
        entities = getattr(data_model, entity_type)
        return find_entity_by_name(entities, target_name)
    except Exception as e:
        print(f"Error adding {entity_type}: {e} target_name:{target_name}")
        return None


def get_group_by_id(groups, target_name, data_model: Data, supabase_client: SupaBaseWorker) -> Group:
    if (
        target_name.replace("_", "-").replace("—", "-").lower()
        == "22пса-1,22пса-2,22пса-3"
    ):
        target_name = "22пса-1"
    return add_and_get_entity(
        entity_type="GROUPS",
        add_func=supabase_client.addGroup,
        entities=groups,
        target_name=target_name.replace("_", "-").replace("—", "-").lower(),
        data_model=data_model,
        args=[],
    )


def get_cabinet_by_id(
        cabinets: list[Cabinet],
        target_name: str,
        data_model: Data,
        supabase_client: SupaBaseWorker,
) -> Cabinet:
    return add_and_get_entity(
        entity_type="CABINETS",
        add_func=supabase_client.addCabinet,
        entities=cabinets,
        target_name=target_name,
        data_model=data_model,
        args=[],
    )


def get_course_by_id(
        courses,
        target_name,
        data_model: Data,
        supabase_client: SupaBaseWorker,
        args: List[str],
) -> Course:
    return add_and_get_entity(
        entity_type="COURSES",
        add_func=supabase_client.addCourse,
        entities=courses,
        target_name=target_name,
        data_model=data_model,
        args=args,
    )


def get_teacher_by_id(
        teachers, target_name, data_model: Data, supabase_client: SupaBaseWorker
) -> Teacher:
    teacher = find_entity_by_name(teachers, target_name)
    if teacher:
        return teacher

    teacher = get_teacher_from_short_name(teachers=teachers, shortName=target_name)
    if teacher:
        return teacher

    return add_and_get_entity(
        entity_type="TEACHERS",
        add_func=supabase_client.addTeacher,
        entities=teachers,
        target_name=target_name,
        data_model=data_model,
        args=[],
    )


def get_teacher_from_short_name(teachers: list[Teacher], shortName: str):
    shortcomparer = (
        shortName.replace(".", "").replace(",", "").replace(" ", "").lower().strip()
    )

    for teacher in teachers:
        fio = teacher.name.split()
        if len(fio) < 3:
            continue

        # Формирование короткой версии имени преподавателя
        compare_result = f"{fio[0]}{fio[1][0]}{fio[2][0]}".lower().strip()

        if (
                compare_result == shortcomparer
                or count_different_characters(compare_result, shortcomparer) == 1
        ):
            print(f"taker {compare_result} and {shortcomparer}")
            print(f"set {teacher}")
            return teacher

    return None


def count_different_characters(str1, str2):
    if len(str1) != len(str2):
        return -1

    # Подсчет количества различных символов
    return sum(char1 != char2 for char1, char2 in zip(str1, str2))
