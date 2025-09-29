from io import BytesIO

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table

from src.alchemy.database_local import (Zamena, ZamenaGroup, ZamenaGroupType,
                                        ZamenaSwaps)
from src.utils.logger import logger


def clean_trash(string: str) -> str:
    return string.replace('(улзвалиди7а)','')

def all_equal(items: list[str]) -> bool:
    return len(set(items)) <= 1


def clean_dirty_string(string: str) -> str:
    return (
        string.replace(" ", "")
        .replace(".", "")
        .replace(",", "")
        .replace("-", "")
        .replace("_", "")
        .replace("\n", " ")
        .replace("\t", "")
        .replace("—", "")
        .replace("—", "")
    ).lower().replace(' ','')


async def parse_zamena_v3(stream: BytesIO, session):
    exceptions: list = []
    zamena_swaps = []
    
    docx: DocumentObject = Document(stream)
    all_rows: list[list[str]] = extract_all_tables_to_rows(docx.tables)
    work_rows: list = all_rows
    # header_paragraphs: List[Paragraph] = docx.paragraphs    
    
    first_row = work_rows[0]
    # # Удаление столбца Время
    # if first_row[0] == 'Время':
    #     for row in work_rows:
    #         row.pop(0)
            
    # Удаление строки хедеров
    if first_row[0] == 'Время':
        work_rows.pop(0)
        
    # Очистка пустых строк
    work_rows = [sublist for sublist in work_rows if any(item != "" for item in sublist)]


    # Простановка группы в первую ячейку
    # ['25ПД-2', '25ПД-2', '25ПД-2', '25ПД-2', '25ПД-2', '25ПД-2', '25ПД-2']
    # ['', '3', '', '', 'Математика', 'Гайсин И.И.', '223']
    # -> 
    # ['25ПД-2', '3', '', '', 'Математика', 'Гайсин И.И.', '223']
    group_name: str = work_rows[0][0]
    for row in work_rows:
        if row[0] == '':
            row[0] = group_name
        else:
            group_name = row[0]


    from src.api_v1.groups.crud import get_groups_normalized_contains

    # Восстановление строк с полной заменой ['','','21П-2','',''] -> ['21П-2','21П-2','21П-2','21П-2','21П-2']
    # for row in work_rows:
    #     non_empty_cells: list[str] = [cell for cell in row if isinstance(cell, str) and cell.strip()]
    #     if not non_empty_cells:
    #         continue
    #     if len(non_empty_cells) == 1:
    #         non_empty_cell: str = clean_dirty_string(non_empty_cells[0])
    #         groups = await get_groups_normalized_contains(session=session, raw_name=non_empty_cell)
    #         if groups:
    #             group = groups[0]
    #             row[:] = [group.name] * len(row)
    #         else:
    #             print(f'🔴 Не найдена группа -> {non_empty_cell}')
    # Восстановление оторванных строк
    # ['4,5', '', '', 'Правовые основы оперативно-', 'Музафаров Ф.Ф.', '112']
    # ['', '', '', 'розыскной \nдеятельности', '', '']
    # -> ['4,5', '', '', 'Правовые основы оперативно-розыскной \nдеятельности', 'Музафаров Ф.Ф.', '112']
    # merged_rows: list[list[str]] = []
    # for row in work_rows:
    #     if row[0] == '' and merged_rows:
    #         prev_row: list[str] = merged_rows[-1]
    #         prev_row[3] = (prev_row[3] + row[3]).strip()
    #     else:
    #         merged_rows.append(row)
    # work_rows = list(merged_rows)
    # перевод пар 3,4 на отдельные строки
    extracted: list = []
    for row in work_rows:
        timings_text = row[1]

        if not timings_text.replace(',','').isdigit():
            extracted.append(row)
            continue
    
        cell: str = timings_text.replace('.', ',')
    
        if cell[0] == ',':
            cell = cell[1:]
        
        if cell[-1] == ',':
            cell = cell[:-1]
        
        timings: list[str] = cell.split(',')
    
        if len(timings) > 1:
            for timing in timings:
                copy_row = row.copy()
                copy_row[1] = timing
                extracted.append(copy_row)
        else:
            extracted.append(row)

    work_rows = list(extracted)
    
    # # Очистка от лишних символов
    work_rows = [[clean_dirty_string(cell) for cell in row] for row in work_rows]
    
    # # Перевод в айдишники
    from src.api_v1.disciplines.crud import \
        find_group_disciplines_by_alias_or_name_or_code_discipline_name
    
    current_group = None
    full_swap_groups_ids = []
    for row in work_rows:
        group_text: str = clean_trash(row[0])
        groups = await get_groups_normalized_contains(
            session=session,
            raw_name=group_text
        )

        if not groups:
            raise Exception(f'🔴 Не найдена группа {group_text}')
        if len(groups) > 1:
            raise Exception(f'🔴 Больше 1 совпадения группы {group_text}')

        group = groups[0]
        # row[:] = [group] * len(row)
        current_group = group

        if all_equal(row):
            full_swap_groups_ids.append(group.id)
            continue
        
        course_text: str = row[4]
        course = None
        if course_text != 'нет':
            founded_disciplines = await find_group_disciplines_by_alias_or_name_or_code_discipline_name(
                session = session,
                group = group,
                raw = course_text
            )
        
            if len(founded_disciplines) == 0:
                exceptions.append(f'🔴 Не найдена дисциплина {course_text} для группы {current_group.name}')
                continue
            if len(founded_disciplines) > 1:
                raise Exception(f'🔴 Больше 1 совпадения дисциплин {course_text} для группы {current_group.name} -> {founded_disciplines}')
            course = founded_disciplines[0]
            
            
        cabinet_text: str = row[6]
        cabinet = None
        if cabinet_text != '':
            from src.parser.schedule.schedule_parser_v3 import \
                find_cabinet_in_cabinets_by_name
            cabinet = await find_cabinet_in_cabinets_by_name(session = session, raw_name = cabinet_text)
            
            if cabinet is None:
                exceptions.append(f'🔴 Не найден кабинет {cabinet_text}')
                
        teacher_text: str = row[5]
        teacher = None
        if teacher != '':
            from src.parser.schedule.schedule_parser_v3 import \
                find_teacher_in_teachers_by_name
            teacher = await find_teacher_in_teachers_by_name(session = session, raw_name = teacher_text)
            
            if teacher is None:
                exceptions.append(f'🔴 Не найден преподаватель {teacher_text}')
        
        timing = int(row[1])
        
        if len(exceptions) != 0:
            continue
        
        # swap = ZamenaSwaps(
        #     group_id = group.id,
        #     timing_id = timing,
        #     teacher_id = teacher.id,
        #     discipline_id = course.id,
        #     cabinet_id = cabinet.id
        # )
        # zamena_swaps.append(swap)
        
    
    # zamena = Zamena(
    #     date_ = ,
    #     saturday_timings = False,
    #     file_url = '',
    #     file_hash = '',   
    # )
    
    # for zamena_swap in zamena_swaps:
    #     zamena_swap.zamena_id = zamena.id
        
    for row in work_rows:
        print(row)
            
    if len(exceptions) > 0:
        for exception in exceptions:
            logger.error(exception)

        raise Exception(exceptions)
    
    
    return {
        'result': 'completed',
        'work_rows': str(work_rows)
    }
    
    
def extract_all_tables_to_rows(tables: list[Table]) -> list[list[str]]:
    rows: list = []
    for table in tables:
        for row in table.rows:
            data: list = []
            for cell in row.cells:
                cell_text: str = cell.text.strip()
                if cell.tables:
                    nested_table_rows: list = extract_all_tables_to_rows(cell.tables)
                    data.extend(nested_table_rows)
                else:
                    data.append(cell_text)
            rows.append(data)
    return rows