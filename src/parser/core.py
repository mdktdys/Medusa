"""
Модуль харнит в себе методы по работе со строкой полученнной из БД
"""

import re
import urllib
from urllib.request import urlopen

import requests
from src.parser.models.zamena_model import Zamena
from src.parser import supabase as supabase_worker
import docx as Docx
from src.models.models import *

from typing import List
from fitz.table import Table
from bs4 import BeautifulSoup
from io import BytesIO
import datetime

from my_secrets import SCHEDULE_URL, BASEURL
from src.parser.models.data_model import Data
from src.parser.models.zamena_table_model import ZamTable


def getParaNameAndTeacher(
    para: str, data_model: Data, supabase_worker: supabase_worker.SupaBaseWorker
) -> None | list[str]:
    if not para:
        return None

    cell_text = para.replace("\n", " ").replace("\t", " ")
    cell_text = re.sub(r" {2,}", " ", cell_text)

    # в ячейке одно слово
    if len(cell_text.strip().split(" ")) == 1:
        print(f"ЧТО ЗА ХУЙНЯ {cell_text}")
        founded_teachers = get_teachers_from_string(
            teachers=data_model.TEACHERS, shortName=cell_text
        )
        if founded_teachers is None:
            return ["", cell_text]

    sample = cell_text.replace(".", "").replace(" ", "").lower()

    founded_teachers = get_teachers_from_string(
        teachers=data_model.TEACHERS, shortName=sample
    )

    if founded_teachers is None:

        founded_raw = get_teacher_from_synonyms_in_raw_text(
            teachers=data_model.TEACHERS, search_text=cell_text
        )
        if founded_raw[0] is not None:
            return [founded_raw[0], cell_text.replace(founded_raw[1], "").strip()]

        search_text: str = cell_text.strip().split(" ")
        for element in search_text:
            founded_course: Course | None = supabase_worker.get_course_from_synonyms(
                element, courses=data_model.COURSES
            )
            if founded_course is not None:
                teacher = cell_text.replace(founded_course.name, "").strip()
                accurate_teacher = supabase_worker.get_teacher_from_synonyms(
                    search_text=teacher, teachers=data_model.TEACHERS
                )
                if accurate_teacher is None:
                    return [teacher, founded_course.name]
                return [accurate_teacher.name, founded_course.name]

    else:
        if len(founded_teachers) == 1:
            return [
                founded_teachers[0].name,
                _clean_teacher_name(sample, founded_teachers[0].name),
            ]

    # if len(founded_teachers) > 1:
    #     return _handle_multiple_teachers(founded_teachers, sample)

    if "Резерв" in cell_text:
        print("see reserveed")
        return _handle_reserved(cell_text)

    return _handle_excpetions_teacher(sample, cell_text)


def _clean_teacher_name(sample: str, teacher_name: str):
    return sample.replace(teacher_name.replace(" ", "").replace(".", "").lower(), "")


def _handle_multiple_teachers(founded_teachers: list[Teacher], sample: str):
    course_text = sample
    for teacher in founded_teachers:
        temp = teacher.name.split(" ")
        short_fio = f"{temp[0]}{temp[1][0]}{temp[2][0]}".lower()
        course_text = course_text.replace(short_fio, "")
    return [founded_teachers[0].name, course_text]


def _handle_reserved(ParaMonday: str):
    parts = ParaMonday.split(" ")
    if len(parts) > 3:
        prepodMonday = f"{parts[-3]} {parts[-2]}"
        remainder = re.sub(r" {2,}", " ", ParaMonday.replace(prepodMonday, "").strip())
        return [prepodMonday, remainder]
    else:
        prepodMonday = parts[-1]
        return [prepodMonday, ParaMonday.replace(prepodMonday, "").strip()]


def _handle_excpetions_teacher(sample, ParaMonday):
    try:
        prepodMonday = f"{sample[-3]} {sample[-2]} {sample[-1]}"
    except IndexError:
        try:
            prepodMonday = f"{sample[-2]} {sample[-1]}"
        except IndexError:
            try:
                prepodMonday = f"{sample[-1]}"
            except IndexError:
                prepodMonday = sample

    res = [prepodMonday.strip(), ParaMonday.replace(prepodMonday, "").strip()]
    print(f"RESULT {res}")
    return res


def removeDoubleRows(table):
    alreadyExist = []
    for row in table:
        if not row in alreadyExist:
            alreadyExist.append(row)
    return alreadyExist


def removeDuplicates(table):
    index = 0
    cleared = []
    for row in table:
        if (table[index - 1])[0] == row[0]:
            continue
        else:
            cleared.append(row)
            index = index + 1
    return cleared


def recoverTeachers(table):
    aww = 0
    # for i in table:
    #     if len(i) < 5 and i[0] != '':
    #         table.remove(i)
    for row in table:
        if row[0] == "":
            if len(row[1].split(" ")) > 2:
                text = table[aww - 1][1]
                table[aww - 1][1] = text + " \n" + row[1]
            if len(row[3].split(" ")) > 2:
                text = table[aww - 1][3]
                table[aww - 1][3] = text + " \n" + row[3]
            if len(row[5].split(" ")) > 2:
                text = table[aww - 1][5]
                table[aww - 1][5] = text + " \n" + row[5]
            if len(row[7].split(" ")) > 2:
                text = table[aww - 1][7]
                table[aww - 1][7] = text + " \n" + row[7]
            if len(row[9].split(" ")) > 2:
                text = table[aww - 1][9]
                table[aww - 1][9] = text + " \n" + row[9]
            if len(row[11].split(" ")) > 2:
                text = table[aww - 1][11]
                table[aww - 1][11] = text + " \n" + row[11]
            table.remove(row)
        else:
            # if()
            pass
        aww += 1
    return table


def defineGroups(groups, table):
    group_index = 0
    groupParas = []
    group = groups[0]
    divided = {}

    if table[0][0] != "№":
        raise Exception("invalid table")

    for row_index in range(1, len(table)):
        # print(group)
        # print(divided.get(group))
        if table[row_index][0] == "№":
            print(f"UPDATE AFTER ROW")
            print(f"{table[row_index - 1]}")
            print(f"{table[row_index]}")
            print(f"{table[row_index + 1]}")
            try:
                divided[group] = groupParas
                group_index = group_index + 1
                group = groups[group_index]
                groupParas = []
            except Exception as err:
                print(err)
                raise Exception(err)
        else:
            groupParas.append(table[row_index])
    else:
        divided[group] = groupParas
        pass
    return divided


def clearTime(rows):
    cleared = []
    for row in rows:
        if row[1] != "":
            row.pop(1)
            cleared.append(row)
        else:
            cleared.append(row)
    return cleared


def clearSingleStrings(table):
    cleared = []
    for row in table:
        if not isinstance(row, str):
            cleared.append(row)
    return cleared


def clearDiscipline(table):
    for row in table:
        if len(row) >= 12:
            if (
                row[0] == ""
                and row[1] == ""
                and row[2] == ""
                and row[3] == ""
                and row[4] == ""
                and row[5] == ""
                and row[6] == ""
            ):
                table.remove(row)
                continue
            if row[1].__contains__("Дисциплина, вид занятия, преподаватель") or row[
                2
            ].__contains__("Дисциплина, вид занятия, преподаватель"):
                table.remove(row)
                continue
            pass
    return table


def get_cabinet_by_id(cabinets: List[Cabinet], target_name: str, sup, data) -> Cabinet:
    for cabinet in cabinets:
        if cabinet.name.lower().replace(" ", "").replace(".", "").replace(
            "-", ""
        ).replace("\n", "") == target_name.lower().replace(" ", "").replace(
            ".", ""
        ).replace(
            "-", ""
        ).replace(
            "\n", ""
        ):
            return cabinet
        else:
            continue
    try:
        supabase_worker.addCabinet(target_name, sup=sup, data=data)
        return get_cabinet_by_id(
            cabinets=data.CABINETS, target_name=target_name, sup=sup, data=data
        )
    except:
        return None
    return None


def count_matching_characters(str1: str, str2: str) -> int:
    return sum(1 for char1, char2 in zip(str1, str2) if char1 == char2)


def count_different_characters(str1, str2):
    if len(str1) != len(str2):
        return -1

    count = 0
    for char1, char2 in zip(str1, str2):
        if char1 != char2:
            count += 1
    return count


def get_teacher_from_synonyms_in_raw_text(teachers: List[Teacher], search_text: str):
    for i in teachers:
        for syn in i.synonyms:
            if search_text.__contains__(syn):
                return (i.name, syn)
    return (None, None)


def get_teachers_from_string(teachers: List[Teacher], shortName: str) -> List[Teacher]:
    short_comparer = (
        shortName.replace(".", "").replace(",", "").replace(" ", "").lower().strip()
    )
    founded = []
    for teacher in teachers:
        fio: List[str] = teacher.name.split(" ")
        if (
            len(fio) > 2
            and fio[0].strip() != ""
            and fio[1].strip() != ""
            and fio[2].strip() != ""
        ):
            if short_comparer.__contains__(
                teacher.name.replace(" ", "").lower().strip()
            ):
                return [teacher]
            compare_result = f"{fio[0]}{fio[1][0]}{fio[2][0]}".lower().strip()
            if short_comparer.__contains__(compare_result):
                founded.append(teacher)

    # if len(founded) > 0:

    return None

    # best_match = None
    # max_matches = -1

    # for teacher in teachers:
    #     fio: List[str] = teacher.name.split(' ')
    #     if len(fio) > 2 and fio[0].strip() != '' and fio[1].strip() != '' and fio[2].strip() != '':
    #         compare_result = f"{fio[0]}{fio[1][0]}{fio[2][0]}".lower().strip()
    #         matches = count_matching_characters(compare_result, short_comparer)

    #         if matches > max_matches:
    #             max_matches = matches
    #             best_match = teacher
    # return [best_match] if best_match else []


def get_teacher_from_short_name(teachers: List[Teacher], shortName: str):
    for i in teachers:
        fio: List[str] = i.name.split(" ")
        if (
            len(fio) > 2
            and fio[0].strip() != ""
            and fio[1].strip() != ""
            and fio[2].strip() != ""
        ):
            compare_result = f"{fio[0]}{fio[1][0]}{fio[2][0]}".lower().strip()
            shortcomparer = (
                shortName.replace(".", "")
                .replace(",", "")
                .replace(" ", "")
                .lower()
                .strip()
            )
            if compare_result == shortcomparer:
                return i
            else:
                if count_different_characters(compare_result, shortcomparer) == 1:
                    print(f"taker {compare_result} and {shortcomparer}")
                    print(f"set {i}")
                    return i
    return None


def get_teacher_by_id(teachers, target_name, sup, data) -> Teacher:
    for teacher in teachers:
        if teacher.name == target_name:
            return teacher
        else:
            search = get_teacher_from_short_name(
                teachers=teachers, shortName=target_name
            )
            if search is not None:
                return search
    print(f"{target_name}")
    # supabase_worker.addTeacher(target_name, sup=sup, data=data)
    raise Exception(target_name)
    return get_teacher_by_id(
        teachers=data.TEACHERS, target_name=target_name, sup=sup, data=data
    )


def get_group_by_id(groups, target_name, supabase_worker, data) -> Group:
    for group in groups:
        if group.name.upper() == target_name.replace("_", "-").upper():
            return group
        else:
            continue
    # supabase_worker.addGroup(target_name.upper(), sup=sup, data=data)
    raise Exception(target_name)
    return get_group_by_id(
        groups=data.GROUPS, target_name=target_name.upper(), sup=sup, data=data
    )
    return None


def get_course_by_id(
    courses, target_name, sup: supabase_worker.SupaBaseWorker, data
) -> Course:
    for course in courses:
        if course.name == target_name.lower():
            return course
        else:
            continue
    founded_course: Course | None = sup.get_course_from_synonyms(
        target_name, courses=data.COURSES
    )
    if founded_course is None:
        raise Exception(target_name)
        # return get_course_by_id(data.COURSES, target_name=target_name, sup=sup, data=data)
        # supabase_worker.addCourse(target_name, sup=sup, data=data)
    return founded_course


def extract_all_tables_to_rows(tables: List[Table]) -> List[List[str]]:
    rows = []
    for table in tables:
        for row in table.rows:
            data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell.tables:
                    nested_table_rows = extract_all_tables_to_rows(cell.tables)
                    data.extend(nested_table_rows)
                else:
                    data.append(cell_text)
            rows.append(data)
    return rows


def is_nested(row: List[str]) -> bool:
    return isinstance(row, List) and all(isinstance(item, str) for item in row)


def remove_duplicates(input_list):
    unique_list = []
    for sublist in input_list:
        if sublist not in unique_list:
            unique_list.append(sublist)
    return unique_list


def check_family(i):
    if (
        (len(i[1].split(" ")) > 2)
        or (len(i[3].split(" ")) > 2)
        or (len(i[5].split(" ")) > 2)
        or (len(i[7].split(" ")) > 2)
        or (len(i[9].split(" ")) > 2)
        or (len(i[11].split(" ")) > 2)
    ):
        if (
            i[2] == ""
            and i[4] == ""
            and i[6] == ""
            and i[8] == ""
            and i[10] == ""
            and i[12] == ""
        ):
            return True
    return False


def remove_headers(rows):
    cleared = []
    for i in rows:
        if i[0] != "Группа":
            cleared.append(i)
    return cleared


def removeDemoExam(rows):
    cleared = []
    for i in rows:
        if not i[0].__contains__("экзамен"):
            cleared.append(i)
    return cleared


def get_all_tables(filename: str):
    doc = Docx.Document(filename)
    tables = doc.tables
    all_rows = extract_all_tables_to_rows(tables)
    return all_rows, doc.paragraphs


def clear_empty_sublists(nested_list: List[List[str]]) -> List[List[str]]:
    return [sublist for sublist in nested_list if any(item != "" for item in sublist)]


def clearNonDataRows(rows: List[List[str]]) -> List[List[str]]:
    filtered_rows = [row for row in rows if len(row) >= 7]
    return filtered_rows


def downloadFile(link: str, filename: str):
    response = requests.get(link)
    if response.status_code == 200:
        with open(filename, "wb") as file:
            file.write(response.content)
        print(f"File '{filename}' has been downloaded successfully.")
    else:
        print("Failed to download the file.")


def getLastZamenaLink(soup: BeautifulSoup):
    # days, month = getMonthAvalibleDays(soup=soup, monthIndex=0)
    # day = days[-1]
    # year = datetime.datetime.now().year.real
    # date = datetime.date(year=year, month=month, day=day)
    table = getAllMonthTables(soup)[0]
    return table.zamenas[-1].link, table.zamenas[-1].date


def getLastZamenaDate(soup: BeautifulSoup):
    days = getMonthAvalibleDays(soup=soup, monthIndex=0)
    month, year = str(getMonthsList(soup=soup)[0]).split(" ")
    date = datetime.date(2025, convertMonthNameToIndex(month) + 1, days[-1])
    return date


def getDaylink(soup: BeautifulSoup, monthIndex: int, day: int):
    table, month = getMonthTable(soup=soup, monthIndex=monthIndex)
    for link in table.find_all("a"):
        text = link.get_text()
        if link:
            if text.isdigit() and int(text) == day:
                return link.get("href")


def convertMonthNameToIndex(name: str):
    months = [
        "январь",
        "февраль",
        "март",
        "апрель",
        "май",
        "июнь",
        "июль",
        "август",
        "сентябрь",
        "октябрь",
        "ноябрь",
        "декабрь",
    ]
    return months.index(name.lower())


def getMonthTable(soup: BeautifulSoup, monthIndex: int):
    newtables = soup.find_all(name="table")
    # newtables = soup.find_all('table', {'class': 'MsoNormalTable'})
    # oldtables = soup.find_all('table', {'class': 'calendar-month'})
    # newtables.extend(oldtables)
    month = (
        convertMonthNameToIndex(
            newtables[monthIndex]
            .find_all("td", {"class": "calendar-month-title"})[0]
            .get_text()
            .split(" ")[0]
        )
        + 1
    )
    tables = getAllMonthTables(soup)
    return newtables[monthIndex], month


def getAllMonthTables(soup: BeautifulSoup) -> List[ZamTable]:
    zam_tables: List[ZamTable] = []
    new_tables = soup.find_all(name="table")
    for i in new_tables:
        class_type = "".join(i["class"]).strip()
        if class_type == "calendar-month":
            if len(i.find_all("td", {"class": "calendar-month-title"})) == 0:
                header = "Ничего"
                if i.find_all("a", {"class": "calendar-month-title"}):
                    header = i.find_next(
                        "a", {"class": "calendar-month-title"}
                    ).get_text()
                else:
                    header = (
                        i.find_previous().find_previous().get_text().replace("\xa0", "")
                    )
                year = int(datetime.datetime.now().year)
                if len(header.split(" ")) > 1:
                    index = convertMonthNameToIndex(header.split(" ")[0])
                    year = int(header.split(" ")[1])
                else:
                    index = convertMonthNameToIndex(header.replace(" ", ""))
                zam_tables.append(ZamTable(raw=i, month_index=index + 1, year=year))
                pass
            else:
                header = (
                    i.find_all("td", {"class": "calendar-month-title"})[0]
                    .get_text()
                    .replace("\xa0", "")
                    .split(" ")
                )
                index = convertMonthNameToIndex(header[0])
                year = int(header[1])
                zam_tables.append(ZamTable(raw=i, month_index=index + 1, year=year))
            pass
        if class_type == "MsoNormalTable":
            header = i.find_next(name="strong").get_text().replace("\xa0", "")
            year = 2025
            index = convertMonthNameToIndex(header)
            zam_tables.append(ZamTable(raw=i, month_index=index + 1, year=year))
            pass
        if class_type == "ui-datepicker-calendar":
            header = i.find_previous().get_text().replace("\xa0", "").split(" ")
            index = convertMonthNameToIndex(header[0])
            year = int(header[1])
            zam_tables.append(ZamTable(raw=i, month_index=index + 1, year=year))
            pass
    return zam_tables


def getAllTablesLinks(tables: List[ZamTable]) -> List[str]:
    links: List[str] = []
    for table in tables:
        links.extend(table.links)
    return links


def get_all_tables_zamenas(tables: List[ZamTable]) -> List[Zamena]:
    zamenas: List[Zamena] = []
    for table in tables:
        zamenas.extend(table.zamenas)
    return zamenas


def getMonthAvalibleDays(soup: BeautifulSoup, monthIndex: int):
    days = []
    table, month = getMonthTable(soup=soup, monthIndex=monthIndex)
    links = table.find_all("a")
    for link in links:
        if link.get_text().isdigit():
            if link:
                days.append(int(link.get_text()))
    return days, month


def getMonthsList(soup: BeautifulSoup):
    paragraphs_with_class = soup.find_all("p", class_="MsoNormal")
    list = []
    for par in paragraphs_with_class:
        paragraphText = par.get_text(strip=True)
        if (
            paragraphText != ""
            and paragraphText is not None
            and not paragraphText.isdigit()
            and paragraphText not in ["Пн", "Вт", "Ср", "Чт", "Пт", "Сб", "Вс"]
        ):
            list.append(par.get_text(strip=True))
    index = 0
    for month in list:
        if len(month.split(" ")) == 1:
            list[index] = f"{month} {list[index + 1].split(' ')[1]}"
        index = index + 1
    return list


def getLatestScheduleMonth(soup: BeautifulSoup) -> str:
    table = soup.find("table", {"class": "ui-datepicker-calendar"})
    paragraphs_with_class = soup.find_all("p", class_="MsoNormal")

    for par in paragraphs_with_class:
        if par.get_text(strip=True) != "" and par.get_text(strip=True) is not None:
            return par.get_text(strip=True).split()[0]


def getLatestSchedleFile():
    html = urlopen(SCHEDULE_URL).read()
    soup = BeautifulSoup(html, "html.parser")

    table = soup.find("table", {"class": "calendar-month"})
    paragraphs_with_class = soup.find_all("p", class_="MsoNormal")

    for par in paragraphs_with_class:
        if par.get_text(strip=True) != "" and par.get_text(strip=True) is not None:
            print(par.get_text(strip=True).split()[0])
            break
    urls = []
    rows = table.find_all("tr")
    for row in rows:
        cells = row.find_all("td")
        for cell in cells:
            link = cell.find("a")
            if link:
                href = link.get("href")
                date = cell.get_text(strip=True)
                if date.isdigit():
                    print(f"Date: {date}, Href: {href}")
                    file_url = urllib.parse.urljoin(BASEURL, href)
                    urls.append(file_url)

    return urls[-1]


def ParasGroupToSoup(
    group, paras, startday, supabase_worker: supabase_worker.SupaBaseWorker, data
):
    date = startday
    supabasePARA = []
    for para in paras:
        number = para[0]
        ParaMonday: str = para[1]
        paraTuesday: str = para[3]
        paraWednesday: str = para[5]
        paraThursday: str = para[7]
        paraFriday: str = para[9]
        paraSaturday: str = para[11]
        days = [
            ParaMonday,
            paraTuesday,
            paraWednesday,
            paraThursday,
            paraFriday,
            paraSaturday,
        ]
        loopindex = 0
        for day in days:
            aww = getParaNameAndTeacher(day, data, supabase_worker=supabase_worker)
            if aww is not None:
                teacher = get_teacher_by_id(
                    target_name=aww[0],
                    teachers=data.TEACHERS,
                    sup=supabase_worker,
                    data=data,
                )
                course = get_course_by_id(
                    target_name=aww[1],
                    courses=data.COURSES,
                    sup=supabase_worker,
                    data=data,
                )
                cabinet = get_cabinet_by_id(
                    target_name=para[2 * (loopindex + 1)],
                    cabinets=data.CABINETS,
                    sup=supabase_worker,
                    data=data,
                )
                if (
                    teacher is not None
                    and course is not None
                    and cabinet is not None
                    and teacher.name != ""
                    and course.name != ""
                ):
                    supabasePARA.append(
                        {
                            "group": group.id,
                            "number": number,
                            "course": course.id,
                            "teacher": teacher.id,
                            "cabinet": cabinet.id,
                            "date": str(date + datetime.timedelta(days=loopindex)),
                        }
                    )
                    pass
            loopindex = loopindex + 1

            pass
    supabase_worker.client.table("Paras").insert(supabasePARA).execute()
    pass


def parseParas(
    date, supabase_worker: supabase_worker.SupaBaseWorker, data, stream: BytesIO
):
    doc = Docx.Document(docx=stream)
    # doc = Docx.Document(docx=f"schedule {date}.docx")
    groups = []
    for i in doc.paragraphs:
        header: str = i.text
        if header.__contains__("Группа - "):
            group_name = header.replace("Группа - ", "").replace(" ", "")
            founded_groups = supabase_worker.get_groups_from_string(
                group_name, data_model=data
            )
            if len(founded_groups) == 0:
                supabase_worker.addGroup(name=group_name, data_model=data)
                gr = group_name
                pass
            else:
                gr = founded_groups[0].name
            groups.append(gr)

    # for i in groups:
    #     if get_group_by_id(target_name=i, sup=sup, groups=data.GROUPS, data=data):
    #         pass

    tables = []
    for i in doc.tables:
        tables.append(i)

    tables = extract_all_tables_to_rows(tables)
    fin = []
    temp = []
    for row in tables:
        try:
            if len(row[0]) > 1:
                for w in row:
                    temp.append(w)
            else:
                temp.append(row)
        except:
            pass
        fin.append(i)

    temp = clearDiscipline(temp)
    temp = clearSingleStrings(temp)

    # temp = clearTime(temp)

    # es = []
    # for i in temp:
    #     if len(i) < 3:
    #         continue
    #     else:
    #         es.append(i)

    # old_row = []
    # for i in es:
    #     if i[0] == '' and i[1] == '' and i[2] == '' and i[3] == '' and i[4] == '' and i[5] == '' and i[6] != '' and i[7] == '' and i[8] == '' and i[9] == '' and i[10] == '' and i[11] == '' and i[12] == '' and i[13] == '':
    #         old_row[5] = f"{old_row[5]} \n{i[6]}"
    #     old_row = i

    # temp = es

    def is_all_empty(my_list):
        for item in my_list:
            if item != "":
                return False
        return True

    def para_count_in_row(row):
        if len(row) == 13:
            count = 0
            for i in range(1, len(row)):
                if i % 2 == 1:
                    if row[i] != "":
                        count = count + 2
            if count % 2 == 0:
                return count / 2
            else:
                print(row)
                raise Exception(f"invalid row para count {row}]")
        else:
            raise Exception(f"invalid row {row}")

    def get_row_text_count_without_number_para(row):
        count = 0
        for i in range(1, len(row)):
            if row[i] != "":
                count = count + 1
        return count

    # Clean doublicate cabinets on new line like ['','123','1234','1234'] to line upper
    cleaned_from_unused_cabinets = []
    iteraion = 0
    for i in temp:
        if i[0] == "" and not is_all_empty(i):
            if para_count_in_row(
                temp[iteraion - 2]
            ) == get_row_text_count_without_number_para(temp[iteraion - 1]):
                temp.remove(temp[iteraion - 1])

                # print(temp[iteraion - 2])
                # print()
                # print(temp[iteraion])
                # print(temp[iteraion + 1])
                # print(20 * "-")
                pass
            else:
                cleaned_from_unused_cabinets.append(i)
        else:
            cleaned_from_unused_cabinets.append(i)

        iteraion = iteraion + 1

    cleaned_from_duplicate_cabinets = []
    for i in range(0, len(temp)):
        if len(temp[i]) < 13 and temp[i][0] != "№":
            print(f"REMOVE \n {temp[i - 1]} \n {temp[i]}")
            # temp.pop(i)
            pass
        else:
            cleaned_from_duplicate_cabinets.append(temp[i])

    temp = cleaned_from_duplicate_cabinets
    # for i in temp:
    #     print(i)
    #     if len(i) == 13 and i[0] != '№' and not is_all_empty(i):
    #         print(20*"_")
    #         print(para_count_in_row(temp[iteraion - 2]))
    #         print(get_row_text_count_without_number_para(temp[iteraion - 1]))
    #         print(i)
    #         print(20 * "_")
    #
    #         if para_count_in_row(temp[iteraion - 2]) == get_row_text_count_without_number_para(temp[iteraion - 1]):
    #             temp.remove(temp[iteraion - 1])
    #
    #             # print(temp[iteraion - 2])
    #             # print()
    #             # print(temp[iteraion])
    #             # print(temp[iteraion + 1])
    #             # print(20 * "-")
    #             pass
    #         else:
    #             cleaned_from_unused_cabinets.append(i)
    #     else:
    #         cleaned_from_unused_cabinets.append(i)
    #
    #     iteraion = iteraion + 1

    # raise Exception("end")

    for iter in range(0, len(temp)):
        if temp[iter][0] == "" and not is_all_empty(temp[iter]):
            if get_row_text_count_without_number_para(temp[iter - 1]) > 1:
                # print(20 * "-")
                # print("WANNA COPY")
                # print(temp[iter - 2])
                # print(temp[iter - 1])
                # print(temp[iter])
                # print(temp[iter + 1])
                # print(20 * "-")

                for current_column in range(0, len(temp[iter])):
                    temp[iter - 1][current_column] = (
                        temp[iter - 1][current_column]
                        + " "
                        + temp[iter][current_column]
                    )

    for row in temp:
        if row[0] == "" and not is_all_empty(row):
            # print(f"to delete {row}")
            temp.remove(row)

    for iter in range(0, len(temp)):
        # print(temp[iter])
        if temp[iter][0] == "" and not is_all_empty(temp[iter]):
            # print(20*"-")
            # print(temp[iter - 2])
            # print(temp[iter-1])
            print(temp[iter])
            # print(temp[iter + 1])
            # print(temp[iter +2 ])
            # print(20 * "-")

    new_list = []
    for i in range(0, len(temp)):
        if temp[i] == temp[-1]:
            new_list.append(temp[i])
            continue
        if is_all_empty(temp[i]):
            continue
        new_list.append(temp[i])

    for i in range(0, len(new_list)):
        if new_list[i][0] == "1":
            if new_list[i - 1][0] == "7":
                new_list.insert(i, ["№", "RECOVERED"])

    # for i in new_list:
    #     print(i)

    divided = defineGroups(groups, new_list)
    for gruppa in divided:
        paras = divided[gruppa]
        # divided[gruppa] = removeDuplicates(paras)
        # paras = divided[gruppa]
        divided[gruppa] = removeDoubleRows(paras)
        # paras = divided[gruppa]
        # divided[gruppa] = recoverTeachers(paras)
        # for i in paras:
        #     if len(i) < 10:
        #         print("GERE")
        #         for s in range(12 - len(i)):
        #             i.append('')
        ParasGroupToSoup(
            group=get_group_by_id(
                target_name=gruppa,
                groups=data.GROUPS,
                supabase_worker=supabase_worker,
                data=data,
            ),
            paras=divided[gruppa],
            supabase_worker=supabase_worker,
            startday=date,
            data=data,
        )
    pass
